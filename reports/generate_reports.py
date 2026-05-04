# -*- coding: utf-8 -*-
"""
Генерация отчётов по лабораторным работам в формате .docx (для сдачи по ГОСТ/методичке).
Запуск: pip install -r requirements.txt && python generate_reports.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


SHOT = "[ВСТАВИТЬ СКРИНШОТ]"


def set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def build_lr1(path: Path) -> None:
    doc = Document()
    set_normal_style(doc)

    add_title(doc, "Отчёт по лабораторной работе №1")
    add_para(doc, "Дисциплина: объектно-ориентированное программирование.")
    add_para(doc, "Тема: привязки данных в WPF и паттерн MVVM.")
    doc.add_paragraph()

    add_heading(doc, "Введение")
    add_para(
        doc,
        "Цель работы — на практике разобрать режимы привязки данных (Binding) в WPF и увидеть разницу "
        "между прямой привязкой элементов друг к другу (через ElementName) и привязкой к свойствам ViewModel. "
        "Дополнительно рассмотрены триггеры в стилях (DataTrigger, Trigger, EventTrigger).",
    )
    add_para(
        doc,
        "Задачи: (1) построить приложение с вкладками TabControl; (2) на каждой вкладке показать два подхода к привязке; "
        "(3) использовать TextBox, Label, Slider, CheckBox; (4) оформить код по MVVM (View / ViewModel / Model).",
    )
    doc.add_paragraph()

    add_heading(doc, "Сравнение реализаций MVVM")
    add_para(
        doc,
        "В репозитории предусмотрены две ветки: первая реализация без сторонних пакетов (интерфейс INotifyPropertyChanged "
        "и класс RelayCommand на базе ICommand), вторая — с пакетом CommunityToolkit.Mvvm (атрибуты ObservableProperty и RelayCommand).",
    )
    add_para(
        doc,
        "Практическая разница для начинающего: в «чистом» варианте больше шаблонного кода (ручные свойства и команды), "
        "но проще объяснить на паре, «что под капотом». Вариант с Toolkit короче и ближе к промышленным проектам, "
        "потому что генератор кода убирает рутину. Архитектурно оба варианта одинаково соответствуют MVVM, "
        "потому что View не содержит бизнес-логики, а состояние экрана живёт в ViewModel.",
    )
    add_para(doc, SHOT)
    doc.add_paragraph()

    add_heading(doc, "Виды привязок")
    for title, body, example in [
        (
            "Default (по умолчанию)",
            "Если Mode не указан, WPF берёт режим по умолчанию для конкретного DependencyProperty. "
            "Для TextBox.Text обычно это TwoWay, поэтому ввод в TextBox обновляет источник.",
            "Пример из проекта: TextBox Text=\"{Binding DemoText}\" на вкладке «По умолчанию» (UserControl DefaultBindingTabView).",
        ),
        (
            "TwoWay",
            "Двусторонняя синхронизация: изменения в UI попадают в источник, а программные изменения источника отражаются в UI.",
            "Пример: TextBox UserName и Slider Volume с Mode=TwoWay на вкладке TwoWayTabView; также прямой вариант через ElementName между двумя TextBox.",
        ),
        (
            "OneTime",
            "Значение передаётся из источника к цели ограниченно (по смыслу — «один раз» при установке привязки), "
            "что удобно для «снимка» или статического отображения стартовых данных.",
            "Пример: Label/TextBox с Mode=OneTime к свойству FrozenLabel; кнопка меняет строку во ViewModel, а OneTime-цели не следуют за обновлениями так же, как при TwoWay.",
        ),
        (
            "OneWay",
            "Только от источника к элементу интерфейса. Подходит для подписей статуса, когда обратная запись из UI не нужна.",
            "Пример: отображение статуса и флага занятости на вкладке OneWayTabView; отдельно прямой OneWay от Slider к TextBlock.",
        ),
    ]:
        add_heading(doc, title, level=2)
        add_para(doc, body)
        add_para(doc, example)
        add_para(doc, SHOT)
        doc.add_paragraph()

    add_heading(doc, "Триггеры")
    add_para(
        doc,
        "DataTrigger реагирует на значение в данных (обычно через привязку DataContext). "
        "Trigger — на изменение свойства самого элемента (например, IsMouseOver). "
        "EventTrigger внутри Style позволяет запускать анимацию по маршрутизируемому событию (например, Click) через BeginStoryboard.",
    )
    add_para(
        doc,
        "Примеры в проекте: DataTrigger для подсветки рамки по bool IsAlert; Trigger для наведения мыши; "
        "EventTrigger для короткой анимации масштаба кнопки на вкладке TriggersTabView.",
    )
    add_para(doc, SHOT)
    doc.add_paragraph()

    add_heading(doc, "Заключение")
    add_para(
        doc,
        "В ходе работы закреплены режимы привязки и их влияние на обновление UI, показана роль ViewModel как источника данных "
        "для представления, а также продемонстрированы типовые триггеры WPF. Полученный проект можно использовать как шаблон для следующих работ.",
    )

    doc.save(path)


def build_lr2(path: Path) -> None:
    doc = Document()
    set_normal_style(doc)

    add_title(doc, "Отчёт по лабораторной работе №2")
    add_para(doc, "Тема: локализация интерфейса WPF без перезапуска приложения.")
    doc.add_paragraph()

    add_heading(doc, "Введение")
    add_para(
        doc,
        "Цель — добавить поддержку русского и английского языков и переключать язык во время работы программы через ComboBox. "
        "Базой служит проект лабораторной работы №1 (MVVM), чтобы не смешивать учебные темы.",
    )
    add_para(doc, SHOT)
    doc.add_paragraph()

    add_heading(doc, "Теория локализации")
    add_para(
        doc,
        "Локализация — это вынесение текстов интерфейса из кода в ресурсы и выбор нужного набора строк в зависимости от культуры (ru-RU, en-US). "
        "В .NET для этого используют CurrentUICulture: при смене культуры приложение должно заново получить строки и обновить привязки "
        "(через INotifyPropertyChanged / DynamicResource / перезагрузку словаря).",
    )
    doc.add_paragraph()

    add_heading(doc, "Способ 1 — RESX в основном проекте (ветка lr2-resx)")
    add_para(
        doc,
        "Строки хранятся в файлах Resources/Strings.resx и Resources/Strings.en-US.resx. "
        "Чтение выполняется через ResourceManager. При смене языка обновляются свойства MainViewModel, к которым привязан заголовок окна и заголовки вкладок.",
    )
    add_para(doc, "Как реализовано: класс ResxLocalization возвращает строку по ключу для CurrentUICulture; MainViewModel вызывает перечитывание после смены выбранного языка.")
    add_para(doc, SHOT)
    doc.add_paragraph()

    add_heading(doc, "Способ 2 — ResourceDictionary в XAML (ветка lr2-xaml-dict)")
    add_para(
        doc,
        "Строки описаны как sys:String в словарях Localization/Lang.ru.xaml и Lang.en.xaml. "
        "При переключении языка соответствующий словарь заменяется в Application.Resources.MergedDictionaries, а элементы окна используют DynamicResource.",
    )
    add_para(
        doc,
        "Как реализовано: класс XamlLanguageSwitcher удаляет предыдущий словарь Lang.* и добавляет новый; интерфейс обновляется без перезапуска.",
    )
    add_para(doc, SHOT)
    doc.add_paragraph()

    add_heading(doc, "Способ 3 — внешняя библиотека классов (ветка lr2-external-lib)")
    add_para(
        doc,
        "RESX перенесён в отдельную сборку OOP_Labo01.Localization.dll, а основное приложение только ссылается на проект и вызывает LocalizedStrings.Get(key). "
        "Такой подход удобен, если строки переиспользуются в нескольких программах.",
    )
    add_para(
        doc,
        "Как реализовано: отдельный .csproj с ресурсами и класс LocalizedStrings; главное приложение меняет культуру и обновляет строки во ViewModel, как в способе 1.",
    )
    add_para(doc, SHOT)
    doc.add_paragraph()

    add_heading(doc, "Заключение")
    add_para(
        doc,
        "Сравнены три практических подхода: RESX в приложении, XAML-словари с DynamicResource и вынос ресурсов во внешнюю DLL. "
        "Для учебного проекта важно понимать, где хранятся строки и какой механизм обновляет UI при смене языка.",
    )

    doc.save(path)


def main() -> None:
    root = Path(__file__).resolve().parent
    out_lr1 = root / "LR1_Report.docx"
    out_lr2 = root / "LR2_Report.docx"
    build_lr1(out_lr1)
    build_lr2(out_lr2)
    print(f"OK: {out_lr1}")
    print(f"OK: {out_lr2}")


if __name__ == "__main__":
    main()
